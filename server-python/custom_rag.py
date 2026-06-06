import csv
import fitz
import logging
import math
import re
import requests
import docx
from collections import Counter
from sentence_transformers import SentenceTransformer
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

class CustomRAGEngine:
    def __init__(self):
        """Initializes caches and keeps startup logging lightweight."""
        logger.info("🧠 Initializing Custom RAG Engine (Zero LangChain)...")
        self.models_cache = {}
        logger.info("✅ Custom Engine Ready!")

    def _get_model(self, model_name: str):
        """Loads a model dynamically or returns from cache if already loaded."""
        if model_name not in self.models_cache:
            logger.info("📥 Downloading/Loading embedding model: %s...", model_name)
            try:
                self.models_cache[model_name] = SentenceTransformer(model_name)
            except Exception as exc:
                raise RuntimeError(f"Failed to load embedding model '{model_name}': {exc}") from exc
        return self.models_cache[model_name]

    # ==========================================
    # 1. RAW DATA EXTRACTION (PDF PARSING)
    # ==========================================
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Opens a PDF and extracts raw text from every page.
        """

        logger.info("Extracting text from PDF: %s", file_path)
        doc = fitz.open(file_path)
        try:
            full_text = "\n".join([page.get_text("text") for page in doc])
        finally:
            doc.close()

        return full_text
    
    def extract_text_from_file(self, file_path: str, filename: str) -> str:
        """Routes a file to the correct text extractor based on its extension."""
        if not filename or "." not in filename:
            raise ValueError("Filename must include a valid extension")
        ext = filename.split('.')[-1].lower()
        
        try:
            if ext == 'pdf':
                return self.extract_text_from_pdf(file_path)
            
            elif ext == 'txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
                    
            elif ext == 'docx':
                doc = docx.Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                
            elif ext == 'csv':
                rows = []
                with open(file_path, newline='', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    for row in reader:
                        rows.append(" | ".join(row))
                return "\n".join(rows)

            else:
                raise ValueError(f"Unsupported file format: {ext}")
        except Exception as e:
            raise Exception(f"Failed to read {ext} file: {str(e)}")
    
    def extract_text_from_url(self, url: str) -> str:
        """Fetches a webpage and extracts clean readable text from it."""
        try:
            logger.info("🌐 Scraping URL: %s...", url)
            # Headers dalna zaroori hai warna kuch websites bot samajh kar block kar deti hain
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(url, headers=headers, timeout=15)
            response.raise_for_status() # Agar 404/500 error aaye toh exception throw karega
            
            # HTML ko parse karo
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Faltu cheezein (scripts, styles, navbars) hata do
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
                
            # Clean text nikal lo
            text = soup.get_text(separator=' ', strip=True)
            
            # Extra spaces ko clean karo
            clean_text = re.sub(r'\s+', ' ', text)
            return clean_text
            
        except Exception as e:
            raise Exception(f"Failed to extract text from URL: {str(e)}")

    # ==========================================
    # 2. CUSTOM CHUNK ALGORITHM
    # ==========================================
    # ---------------------------------------------------------
    # CHUNKING STRATEGY 1: NAIVE SLIDING WINDOW (Legacy)
    # ---------------------------------------------------------
    def chunk_text_naive(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> list:
        """
        Splits a large string into smaller chunks with a sliding window overlap.
        We do this manually using string slicing.
        """
        logger.info("Chunking text (Size: %s, Overlap: %s)...", chunk_size, overlap)
        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            #Defining the end of the current chunk
            end = start + chunk_size

            #Extract the substring
            chunk = text[start:end]
            chunks.append(chunk)

            #Move the start pointer forward, but step back by the iverlap amount
            start += (chunk_size - overlap)
        
        return [c for c in chunks if c]
    
    # ---------------------------------------------------------
    # CHUNKING STRATEGY 2: SENTENCE WINDOW (Semantic)
    # ---------------------------------------------------------
    def chunk_text_sentence(self, text: str, sentences_per_chunk: int = 6, overlap: int = 2) -> list:
        """
        Splits text safely by punctuation (., !, ?) so sentences are never cut in half.
        Groups them into windows (e.g., 6 sentences per chunk, 2 overlap).
        """
        # Regex magic: Split by period/exclamation/question mark followed by a space
        sentences = re.split(r'(?<=[.!?])\s+', text)
        sentences = [s.strip() for s in sentences if s.strip()]

        chunks = []
        i = 0
        while i < len(sentences):
            # Join the next 'N' sentences together
            chunk = " ".join(sentences[i : i + sentences_per_chunk])
            chunks.append(chunk)
            # Move forward, but keep 'overlap' number of sentences
            i += (sentences_per_chunk - overlap)
        
        return [c for c in chunks if len(c) > 10]
    
    # ---------------------------------------------------------
    # CHUNKING STRATEGY 3: PARAGRAPH / RECURSIVE
    # ---------------------------------------------------------
    def chunk_text_paragraph(self, text: str, max_length: int = 1200) -> list: # FIXED NAME & TYPE
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""

        for para in paragraphs:
            para = para.strip()
            if not para: continue

            if len(current_chunk) + len(para) > max_length and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = para
            else:
                current_chunk += "\n\n" + para if current_chunk else para
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks # FIXED: Return missing tha!

    # ==========================================
    # 3. DIRECT VECTORIZATION (EMBEDDINGS)
    # ==========================================
    # ---------------------------------------------------------
    # 1. VECTOR SEARCH (Meaning)
    # ---------------------------------------------------------
    def vectorize(self, chunks: list, model_name: str = "all-MiniLM-L6-v2") -> list:
        """
        Converts text chunks into mathematical vectors (arrays of floats) 
        using a raw HuggingFace model.
        """
        logger.info("Generating vectors for %s chunks...", len(chunks))
        # encode() returns a Numpy array. We convert it to a standard Python list of floats
        # because databases like Supabase (pgvector) expect standard arrays/lists.
        model = self._get_model(model_name)
        embeddings = model.encode(chunks).tolist()
        
        return embeddings


    # ==========================================
    # 4. THE MATH: COSINE SIMILARITY FROM SCRATCH
    # ==========================================
    def _cosine_similarity(self, v1: list, v2: list) -> float:
        """
        Calculates the cosine similarity between two vectors purely in Python.
        Returns a value between -1 and 1. (Closer to 1 means highly similar).
        """
        if len(v1) != len(v2):
            raise ValueError("Vectors must be of the same length")
        
        # Step A: Calculate Dot Product (A . B)
        # Multiply corresponding numbers and sum them up
        dot_product = sum(a * b for a, b in zip(v1, v2))

        # Step B: Calculate Magnitude (Length) of both vectors (||A|| and ||B||)
        # Square each number, sum them, and take the square root
        magnitude_v1 = math.sqrt(sum(a * a for a in v1))
        magnitude_v2 = math.sqrt(sum(b * b for b in v2))

        #Step C: Prevent divison by zero
        if magnitude_v1 == 0 or magnitude_v2 == 0:
            return 0.0
        
        #Step D: Applt the formula
        similarity = dot_product / (magnitude_v1 * magnitude_v2)

        return similarity

    # ==========================================
    # 5. THE SEARCH ENGINE
    # ==========================================
    def search(self, query_vector: list, document_vectors: list, top_k: int = 7) -> list:
        """
        Compares the query against ALL documents and returns the top matches.
        """
        logger.info("Searching through %s vectors...", len(document_vectors))

        results = []

        for index, doc_vec in enumerate(document_vectors):
            score = self._cosine_similarity(query_vector, doc_vec)
            results.append({
                "chunk_index": index,
                "similarity_score": score
            })
        
        #Sort results by highest score first (descending)
        results.sort(key=lambda x: x["similarity_score"], reverse=True)

        #Return only thr top K results
        return results[:top_k]

    # ---------------------------------------------------------
    # 2. KEYWORD SEARCH (Exact Match - Custom TF-IDF)
    # ---------------------------------------------------------
    def _tokenize(self, text: str) -> list:
        # Convert to lowercase and extract only words (removes punctuation)
        return re.findall(r'\b\w+\b', text.lower())
    
    def _keyword_search(self, query: str, document_texts: list) -> list:
        query_words = self._tokenize(query)
        N = len(document_texts)

        if not query_words or N == 0:
            return [{"chunk_index": i, "score": 0.0} for i in range(N)]
        
        #Calculate Document Frequency (DF)
        df = Counter()
        doc_tokens_list = [self._tokenize(doc) for doc in document_texts]
        for tokens in doc_tokens_list:
            unique_tokens = set(tokens)
            for w in query_words:
                if w in unique_tokens:
                    df[w] += 1
        
        # Calculate Inverse Document Frequency (IDF)
        idf = {}
        for w in query_words:
            # Add 1 to avoid division by zero (Smooth IDF)
            idf[w] = math.log((N + 1) / (df[w] + 1)) + 1 

        results = []
        for index, tokens in enumerate(doc_tokens_list):
            score = 0.0
            if not tokens:
                results.append({"chunk_index": index, "score": 0.0})
                continue
            
            # Calculate Term Frequency (TF)
            tf = Counter(tokens)
            doc_len = len(tokens)
            for w in query_words:
                term_freq = tf[w] / doc_len
                score += term_freq * idf[w] # Final TF-IDF Score
                
            results.append({"chunk_index": index, "score": score})
            
        return results
    
    # ---------------------------------------------------------
    # 3. HYBRID SEARCH ROUTER (Combining Both)
    # ---------------------------------------------------------
    def hybrid_search(self, query_text: str, query_vector: list, document_texts: list, document_vectors: list, alpha: float = 0.7, top_k: int = 10):
        # 1. FIXED: Correct function names from your custom_rag.py
        # Your code has 'search' instead of 'vector_search'
        vector_results = self.search(query_vector, document_vectors, top_k=max(len(document_texts), 1))
        
        # Your code has '_keyword_search'
        keyword_results = self._keyword_search(query_text, document_texts) 
        
        combined_scores = {}
        
        # 2. Vector Results combine (Note: Using 'similarity_score' as per your search function)
        for r in vector_results:
            v_score = r.get("similarity_score", 0.0) 
            combined_scores[r["chunk_index"]] = v_score * alpha
            
        # 3. Keyword Results combine (Note: Using 'score' as per your _keyword_search function)
        for r in keyword_results:
            k_score = r.get("score", 0.0)
            
            if r["chunk_index"] in combined_scores:
                combined_scores[r["chunk_index"]] += k_score * (1.0 - alpha)
            else:
                combined_scores[r["chunk_index"]] = k_score * (1.0 - alpha)
                    
        # 4. Sort and return top_k
        sorted_results = sorted(
            [{"chunk_index": k, "score": v} for k, v in combined_scores.items()],
            key=lambda x: x["score"],
            reverse=True
        )
        
        return sorted_results[:top_k]

# import fitz
# import math
# import re
# import requests
# import docx
# from collections import Counter
# from sentence_transformers import SentenceTransformer
# from bs4 import BeautifulSoup

# class CustomRAGEngine:
#     def __init__(self):
#         print("🧠 Initializing Custom RAG Engine (Zero LangChain)...")
#         self.models_cache = {}
#         print("✅ Custom Engine Ready!")

#     def _get_model(self, model_name: str):
#         """Loads a model dynamically or returns from cache if already loaded."""
#         if model_name not in self.models_cache:
#             print(f"📥 Downloading/Loading embedding model: {model_name}...")
#             self.models_cache[model_name] = SentenceTransformer(model_name)
#         return self.models_cache[model_name]

#     # ==========================================
#     # 1. RAW DATA EXTRACTION (PDF PARSING)
#     # ==========================================
#     def extract_text_from_pdf(self, file_path: str) -> str:
#         """
#         Opens a PDF and extracts raw text from every page.
#         """

#         print(f"Extracting text from PDF: {file_path}")
#         doc = fitz.open(file_path)
#         full_text = "\n".join([page.get_text("text") for page in doc])

#         # for pag_num in range(len(doc)):
#         #     page = doc.load_page(pag_num)
#         #     full_text += page.get_text("text") + "\n"
        
#         return full_text
    
#     def extract_text_from_file(self, file_path: str, filename: str) -> str:
#         """Dynamically routes the file to the correct text extractor based on extension."""
#         ext = filename.split('.')[-1].lower()
        
#         try:
#             if ext == 'pdf':
#                 return self.extract_text_from_pdf(file_path)
            
#             elif ext == 'txt':
#                 with open(file_path, 'r', encoding='utf-8') as f:
#                     return f.read()
                    
#             elif ext == 'docx':
#                 doc = docx.Document(file_path)
#                 return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                
#             elif ext == 'csv':
#                 text = ""
#                 with open(file_path, newline='', encoding='utf-8') as f:
#                     reader = csv.reader(f)
#                     for row in reader:
#                         text += " | ".join(row) + "\n"
#                 return text
                
#             else:
#                 raise ValueError(f"Unsupported file format: {ext}")
#         except Exception as e:
#             raise Exception(f"Failed to read {ext} file: {str(e)}")
    
#     def extract_text_from_url(self, url: str) -> str:
#         """Fetches a webpage and extracts clean readable text from it."""
#         try:
#             print(f"🌐 Scraping URL: {url}...")
#             # Headers dalna zaroori hai warna kuch websites bot samajh kar block kar deti hain
#             headers = {
#                 'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
#             }
#             response = requests.get(url, headers=headers, timeout=10)
#             response.raise_for_status() # Agar 404/500 error aaye toh exception throw karega
            
#             # HTML ko parse karo
#             soup = BeautifulSoup(response.text, 'html.parser')
            
#             # Faltu cheezein (scripts, styles, navbars) hata do
#             for script in soup(["script", "style", "nav", "footer", "header"]):
#                 script.decompose()
                
#             # Clean text nikal lo
#             text = soup.get_text(separator=' ', strip=True)
            
#             # Extra spaces ko clean karo
#             clean_text = re.sub(r'\s+', ' ', text)
#             return clean_text
            
#         except Exception as e:
#             raise Exception(f"Failed to extract text from URL: {str(e)}")

#     # ==========================================
#     # 2. CUSTOM CHUNK ALGORITHM
#     # ==========================================
#     # ---------------------------------------------------------
#     # CHUNKING STRATEGY 1: NAIVE SLIDING WINDOW (Legacy)
#     # ---------------------------------------------------------
#     def chunk_text_naive(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> list:
#         """
#         Splits a large string into smaller chunks with a sliding window overlap.
#         We do this manually using string slicing.
#         """
#         print(f"Chunking text (Size: {chunk_size}, Overlap: {overlap})...")
#         chunks = []
#         start = 0
#         text_length = len(text)

#         while start < text_length:
#             #Defining the end of the current chunk
#             end = start + chunk_size

#             #Extract the substring
#             chunk = text[start:end]
#             chunks.append(chunk)

#             #Move the start pointer forward, but step back by the iverlap amount
#             start += (chunk_size - overlap)
        
#         return [c for c in chunks if c]
    
#     # ---------------------------------------------------------
#     # CHUNKING STRATEGY 2: SENTENCE WINDOW (Semantic)
#     # ---------------------------------------------------------
#     def chunk_text_sentence(self, text: str, sentences_per_chunk: int = 6, overlap: int = 2) -> list:
#         """
#         Splits text safely by punctuation (., !, ?) so sentences are never cut in half.
#         Groups them into windows (e.g., 6 sentences per chunk, 2 overlap).
#         """
#         # Regex magic: Split by period/exclamation/question mark followed by a space
#         sentences = re.split(r'(?<=[.!?])\s+', text)
#         sentences = [s.strip() for s in sentences if s.strip()]

#         chunks = []
#         i = 0
#         while i < len(sentences):
#             # Join the next 'N' sentences together
#             chunk = " ".join(sentences[i : i + sentences_per_chunk])
#             chunks.append(chunk)
#             # Move forward, but keep 'overlap' number of sentences
#             i += (sentences_per_chunk - overlap)
        
#         return [c for c in chunks if len(c) > 10]
    
#     # ---------------------------------------------------------
#     # CHUNKING STRATEGY 3: PARAGRAPH / RECURSIVE
#     # ---------------------------------------------------------
#     def chunk_text_paragraph(self, text: str, max_length: int = 1200) -> list: # FIXED NAME & TYPE
#         paragraphs = text.split('\n\n')
#         chunks = []
#         current_chunk = ""

#         for para in paragraphs:
#             para = para.strip()
#             if not para: continue

#             if len(current_chunk) + len(para) > max_length and current_chunk:
#                 chunks.append(current_chunk.strip())
#                 current_chunk = para
#             else:
#                 current_chunk += "\n\n" + para if current_chunk else para
        
#         if current_chunk:
#             chunks.append(current_chunk.strip())
        
#         return chunks # FIXED: Return missing tha!

#     # ==========================================
#     # 3. DIRECT VECTORIZATION (EMBEDDINGS)
#     # ==========================================
#     # ---------------------------------------------------------
#     # 1. VECTOR SEARCH (Meaning)
#     # ---------------------------------------------------------
#     def vectorize(self, chunks: list, model_name: str = "all-MiniLM-L6-v2") -> list:
#         """
#         Converts text chunks into mathematical vectors (arrays of floats) 
#         using a raw HuggingFace model.
#         """
#         print(f"Generating vectors for {len(chunks)} chunks...")
#         # encode() returns a Numpy array. We convert it to a standard Python list of floats
#         # because databases like Supabase (pgvector) expect standard arrays/lists.
#         model = self._get_model(model_name)
#         embeddings = model.encode(chunks).tolist()
        
#         return embeddings


#     # ==========================================
#     # 4. THE MATH: COSINE SIMILARITY FROM SCRATCH
#     # ==========================================
#     def _cosine_similarity(self, v1: list, v2: list) -> float:
#         """
#         Calculates the cosine similarity between two vectors purely in Python.
#         Returns a value between -1 and 1. (Closer to 1 means highly similar).
#         """
#         if len(v1) != len(v2):
#             raise ValueError("Vectors must be of the same length")
        
#         # Step A: Calculate Dot Product (A . B)
#         # Multiply corresponding numbers and sum them up
#         dot_product = sum(a * b for a, b in zip(v1, v2))

#         # Step B: Calculate Magnitude (Length) of both vectors (||A|| and ||B||)
#         # Square each number, sum them, and take the square root
#         magnitude_v1 = math.sqrt(sum(a * a for a in v1))
#         magnitude_v2 = math.sqrt(sum(b * b for b in v2))

#         #Step C: Prevent divison by zero
#         if magnitude_v1 == 0 or magnitude_v2 == 0:
#             return 0.0
        
#         #Step D: Applt the formula
#         similarity = dot_product / (magnitude_v1 * magnitude_v1)

#         return similarity

#     # ==========================================
#     # 5. THE SEARCH ENGINE
#     # ==========================================
#     def search(self, query_vector: list, document_vectors: list, top_k: int = 7) -> list:
#         """
#         Compares the query against ALL documents and returns the top matches.
#         """
#         print(f"Searching through {len(document_vectors)} vectors...")

#         results = []

#         for index, doc_vec in enumerate(document_vectors):
#             score = self._cosine_similarity(query_vector, doc_vec)
#             results.append({
#                 "chunk_index": index,
#                 "similarity_score": score
#             })
        
#         #Sort results by highest score first (descending)
#         results.sort(key=lambda x: x["similarity_score"], reverse=True)

#         #Return only thr top K results
#         return results[:top_k]

#     # ---------------------------------------------------------
#     # 2. KEYWORD SEARCH (Exact Match - Custom TF-IDF)
#     # ---------------------------------------------------------
#     def _tokenize(self, text: str) -> list:
#         # Convert to lowercase and extract only words (removes punctuation)
#         return re.findall(r'\b\w+\b', text.lower())
    
#     def _keyword_search(self, query: str, document_texts: list) -> list:
#         query_words = self._tokenize(query)
#         N = len(document_texts)

#         if not query_words or N == 0:
#             return [{"chunk_index": i, "score": 0.0} for i in range(N)]
        
#         #Calculate Document Frequency (DF)
#         df = Counter()
#         doc_tokens_list = [self._tokenize(doc) for doc in document_texts]
#         for tokens in doc_tokens_list:
#             unique_tokens = set(tokens)
#             for w in query_words:
#                 if w in unique_tokens:
#                     df[w] += 1
        
#         # Calculate Inverse Document Frequency (IDF)
#         idf = {}
#         for w in query_words:
#             # Add 1 to avoid division by zero (Smooth IDF)
#             idf[w] = math.log((N + 1) / (df[w] + 1)) + 1 

#         results = []
#         for index, tokens in enumerate(doc_tokens_list):
#             score = 0.0
#             if not tokens:
#                 results.append({"chunk_index": index, "score": 0.0})
#                 continue
            
#             # Calculate Term Frequency (TF)
#             tf = Counter(tokens)
#             doc_len = len(tokens)
#             for w in query_words:
#                 term_freq = tf[w] / doc_len
#                 score += term_freq * idf[w] # Final TF-IDF Score
                
#             results.append({"chunk_index": index, "score": score})
            
#         return results
    
#     # ---------------------------------------------------------
#     # 3. HYBRID SEARCH ROUTER (Combining Both)
#     # ---------------------------------------------------------
#     def hybrid_search(self, query_text: str, query_vector: list, document_texts: list, document_vectors: list, alpha: float = 0.7, top_k: int = 10):
#         # 1. FIXED: Correct function names from your custom_rag.py
#         # Your code has 'search' instead of 'vector_search'
#         vector_results = self.search(query_vector, document_vectors, top_k=len(document_texts))
        
#         # Your code has '_keyword_search'
#         keyword_results = self._keyword_search(query_text, document_texts) 
        
#         combined_scores = {}
        
#         # 2. Vector Results combine (Note: Using 'similarity_score' as per your search function)
#         for r in vector_results:
#             v_score = r.get("similarity_score", 0.0) 
#             combined_scores[r["chunk_index"]] = v_score * alpha
            
#         # 3. Keyword Results combine (Note: Using 'score' as per your _keyword_search function)
#         for r in keyword_results:
#             k_score = r.get("score", 0.0)
            
#             if r["chunk_index"] in combined_scores:
#                 combined_scores[r["chunk_index"]] += k_score * (1.0 - alpha)
#             else:
#                 combined_scores[r["chunk_index"]] = k_score * (1.0 - alpha)
                    
#         # 4. Sort and return top_k
#         sorted_results = sorted(
#             [{"chunk_index": k, "score": v} for k, v in combined_scores.items()],
#             key=lambda x: x["score"],
#             reverse=True
#         )
        
#         return sorted_results[:top_k]